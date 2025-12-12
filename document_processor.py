"""Document processing module for extracting text from various file formats."""

import os
from typing import Optional, Dict, Any
import PyPDF2
from docx import Document
import pandas as pd
from pathlib import Path


class DocumentProcessor:
    """Handles extraction of text from various document formats."""

    @staticmethod
    def process_file(file_path: str) -> Dict[str, Any]:
        """
        Process a file and extract its content.

        Args:
            file_path: Path to the file to process

        Returns:
            Dict containing:
                - text: Extracted text content
                - metadata: File metadata (name, size, type)
                - success: Boolean indicating success
                - error: Error message if failed
        """
        result = {
            "text": "",
            "metadata": {},
            "success": False,
            "error": None
        }

        try:
            file_path = Path(file_path)
            extension = file_path.suffix.lower()

            # Get file metadata
            result["metadata"] = {
                "name": file_path.name,
                "size": os.path.getsize(file_path),
                "type": extension[1:] if extension else "unknown"
            }

            # Process based on file type
            if extension == ".pdf":
                result["text"] = DocumentProcessor._process_pdf(file_path)
            elif extension in [".docx", ".doc"]:
                result["text"] = DocumentProcessor._process_docx(file_path)
            elif extension == ".txt":
                result["text"] = DocumentProcessor._process_txt(file_path)
            elif extension == ".csv":
                result["text"] = DocumentProcessor._process_csv(file_path)
            elif extension in [".xlsx", ".xls"]:
                result["text"] = DocumentProcessor._process_excel(file_path)
            else:
                result["error"] = f"Unsupported file type: {extension}"
                return result

            result["success"] = True

        except Exception as e:
            result["error"] = f"Error processing file: {str(e)}"

        return result

    @staticmethod
    def _process_pdf(file_path: Path) -> str:
        """Extract text from PDF file."""
        text = []
        with open(file_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text.append(f"--- Page {page_num + 1} ---\n{page_text}")
        return "\n\n".join(text)

    @staticmethod
    def _process_docx(file_path: Path) -> str:
        """Extract text from DOCX file."""
        doc = Document(file_path)
        text = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                text.append(" | ".join(row_data))

        return "\n".join(text)

    @staticmethod
    def _process_txt(file_path: Path) -> str:
        """Extract text from TXT file."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
            return file.read()

    @staticmethod
    def _process_csv(file_path: Path) -> str:
        """Extract text from CSV file."""
        df = pd.read_csv(file_path)
        # Convert DataFrame to readable text format
        text = [f"CSV File with {len(df)} rows and {len(df.columns)} columns\n"]
        text.append(f"Columns: {', '.join(df.columns)}\n")
        text.append("\nData Preview:\n")
        text.append(df.to_string(index=False, max_rows=100))

        return "\n".join(text)

    @staticmethod
    def _process_excel(file_path: Path) -> str:
        """Extract text from Excel file."""
        text = []
        excel_file = pd.ExcelFile(file_path)

        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(excel_file, sheet_name=sheet_name)
            text.append(f"\n--- Sheet: {sheet_name} ---")
            text.append(f"{len(df)} rows and {len(df.columns)} columns")
            text.append(f"Columns: {', '.join(df.columns)}")
            text.append("\nData Preview:")
            text.append(df.to_string(index=False, max_rows=100))

        return "\n".join(text)

    @staticmethod
    def extract_financial_info(text: str) -> Dict[str, Any]:
        """
        Extract basic financial information from text.
        This is a simple pattern-based extraction.

        Args:
            text: Text content to analyze

        Returns:
            Dict with potential financial data found
        """
        import re

        result = {
            "amounts": [],
            "dates": [],
            "potential_transactions": []
        }

        # Find currency amounts
        amount_pattern = r'\$?\s*\d{1,3}(?:,\d{3})*(?:\.\d{2})?'
        amounts = re.findall(amount_pattern, text)
        result["amounts"] = [a.strip() for a in amounts if a.strip()]

        # Find dates
        date_patterns = [
            r'\d{1,2}/\d{1,2}/\d{2,4}',
            r'\d{4}-\d{2}-\d{2}',
            r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{1,2},? \d{4}'
        ]
        for pattern in date_patterns:
            dates = re.findall(pattern, text, re.IGNORECASE)
            result["dates"].extend(dates)

        return result
